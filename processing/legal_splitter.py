import re
from docx import Document
from langchain_core.documents import Document as LCDocument


def extract_docx_paragraphs(docx_path: str, skip_texts: set[str] = None) -> list[dict]:
    skip_texts = skip_texts or set()
    doc = Document(docx_path)
    elements = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text or text in skip_texts:
            continue
        is_bold = any(run.bold for run in para.runs if run.text.strip())
        elements.append({"text": text, "bold": is_bold})
    return elements


def split_docx_legal(elements: list[dict], source_name: str,
                      stop_at_patterns: list[str] = None) -> list[LCDocument]:
    stop_at_patterns = stop_at_patterns or []
    chunks = []
    current_chuong = None
    current_muc = None
    current_dieu = None
    current_dieu_title = ""
    current_khoan = None
    current_khoan_opening = ""
    current_khoan_has_diem = False  # MỚI: đánh dấu khoản hiện tại đã sinh chunk điểm (a,b,c..) hay chưa
    buffer = []
    dieu_body = []
    pending_dieu_number = None
    stop_processing = False

    def flush_khoan_as_chunk():
        # SỬA: nếu khoản này đã có các điểm a,b,c... thì nội dung câu mở đầu đã được
        # nhúng vào từng chunk điểm rồi (xem diem_match bên dưới) -- không được tạo
        # thêm 1 chunk "khoản mở đầu" đứng riêng nữa, đó chính là chunk thừa bị lặp lại.
        if buffer and current_dieu and not current_khoan_has_diem:
            content = f"{current_dieu_title}\n{current_khoan_opening}".strip()
            chunks.append(LCDocument(
                page_content=content,
                metadata={"source": source_name, "chuong": current_chuong, "muc": current_muc,
                          "dieu": current_dieu, "khoan": current_khoan, "diem": None}
            ))
        buffer.clear()  # SỬA: clear vô điều kiện, không phụ thuộc current_dieu

    def flush_dieu_body():
        if dieu_body and current_dieu:
            content = f"{current_dieu_title}\n" + "\n".join(dieu_body)
            chunks.append(LCDocument(
                page_content=content.strip(),
                metadata={"source": source_name, "chuong": current_chuong, "muc": current_muc,
                          "dieu": current_dieu, "khoan": None, "diem": None}
            ))
        dieu_body.clear()  # SỬA: clear vô điều kiện

    def reset_state_on_heading():
        """MỚI: gọi khi gặp Chương/Mục -- xóa sạch mọi state dở dang, tránh rò rỉ sang Điều sau."""
        nonlocal current_dieu, current_khoan, current_khoan_opening, current_dieu_title, current_khoan_has_diem
        current_dieu = None
        current_khoan = None
        current_khoan_opening = ""
        current_dieu_title = ""
        current_khoan_has_diem = False
        buffer.clear()
        dieu_body.clear()

    for el in elements:
        text = el["text"]

        if any(re.match(p, text, re.IGNORECASE) for p in stop_at_patterns):
            flush_khoan_as_chunk()
            flush_dieu_body()
            stop_processing = True
            continue
        if stop_processing:
            continue

        chuong_match = re.match(r"CHƯƠNG\s+([IVXLCDM]+)", text, re.IGNORECASE)
        if chuong_match:
            flush_khoan_as_chunk()
            flush_dieu_body()
            reset_state_on_heading()          # SỬA: reset đầy đủ
            current_chuong = chuong_match.group(1)
            current_muc = None
            continue

        muc_match = re.match(r"Mục\s+(\d+)", text, re.IGNORECASE)
        if muc_match:
            flush_khoan_as_chunk()
            flush_dieu_body()
            reset_state_on_heading()          # SỬA: reset đầy đủ
            current_muc = muc_match.group(1)
            continue

        dieu_match = re.match(r"Điều\s+(\d+)\.\s*(.*)", text)
        if dieu_match:
            flush_khoan_as_chunk()
            flush_dieu_body()
            current_dieu = dieu_match.group(1)
            title_rest = dieu_match.group(2).strip()
            if title_rest:
                current_dieu_title = text
                pending_dieu_number = None
            else:
                pending_dieu_number = current_dieu
                current_dieu_title = text
            current_khoan = None
            current_khoan_opening = ""
            current_khoan_has_diem = False
            continue

        if pending_dieu_number:
            current_dieu_title += " " + text
            pending_dieu_number = None
            continue

        khoan_match = re.match(r"(\d+)\.\s", text)
        if khoan_match and current_dieu:
            flush_khoan_as_chunk()
            current_khoan = khoan_match.group(1)
            current_khoan_opening = text
            current_khoan_has_diem = False
            buffer.append(text)
            continue

        diem_match = re.match(r"([a-zđê])\)\s", text)
        if diem_match and current_khoan:
            diem_ky_hieu = diem_match.group(1)
            current_khoan_has_diem = True  # SỬA: đánh dấu để flush_khoan_as_chunk() không tạo chunk thừa nữa
            content = f"{current_dieu_title}\n{current_khoan_opening}\n{text}".strip()
            chunks.append(LCDocument(
                page_content=content,
                metadata={"source": source_name, "chuong": current_chuong, "muc": current_muc,
                          "dieu": current_dieu, "khoan": current_khoan, "diem": diem_ky_hieu}
            ))
            continue


        if current_khoan and current_dieu:
            buffer.append(text)
            current_khoan_opening += "\n" + text
        elif current_dieu:
            dieu_body.append(text)

    flush_khoan_as_chunk()
    flush_dieu_body()
    return chunks


def print_debug_report(chunks: list[LCDocument]):
    """Hàm kiểm tra dùng chung cho mọi file """
    print(f"Tổng số chunk: {len(chunks)}")
    no_dieu = sum(1 for c in chunks if c.metadata.get("dieu") is None)
    print(f"Chunk không xác định Điều: {no_dieu}/{len(chunks)}")

    if no_dieu < len(chunks):
        dieu_so = sorted({int(c.metadata["dieu"]) for c in chunks if c.metadata.get("dieu")}, key=int)
        print(f"Số Điều nhỏ nhất: {dieu_so[0]}, lớn nhất: {dieu_so[-1]}")
        missing = [i for i in range(dieu_so[0], dieu_so[-1] + 1) if i not in dieu_so]
        print(f"Các số Điều bị thiếu: {missing}")

    print("\n--- 30 chunk đầu ---")
    for c in chunks[:30]:
        print("---")
        print(c.page_content)
        print(c.metadata)


